"""Genere la version .docx du CV a partir de cv/cv.html.

    python scripts/make_cv_docx.py

Pourquoi une version Word alors que le PDF est deja propre : les mesures
publiees en 2026 donnent le .docx a plus de 95 % de parsing correct sur
tous les ATS majeurs, contre 83 % pour le PDF sur Taleo — l'un des plus
rigides et des plus repandus. Douze points d'ecart sur un seul moteur.

Le contenu n'est PAS ressaisi ici : il est lu dans cv/cv.html, qui reste
l'unique source. C'est exactement l'erreur que cv.tex a commise en
divergeant silencieusement pendant des mois.

Sur les polices. Le PDF embarque Newsreader et IBM Plex Sans, donc il
s'affiche partout a l'identique. Un .docx ne peut pas offrir ca : Word
utilise les polices installees chez le lecteur, et substitue le reste.
On choisit donc des polices presentes sur toutes les machines et qui
tiennent les memes roles : Georgia pour le serif des intitules, Arial
pour le texte. Le document ne sera pas identique au PDF, il en sera la
traduction la plus proche possible.

Regles ATS specifiques au format Word :
  - coordonnees dans le CORPS du document, jamais dans un en-tete Word :
    la majorite des ATS ignorent en-tetes et pieds de page, ce qui ferait
    disparaitre nom, e-mail et telephone ;
  - aucun tableau, aucune zone de texte, une seule colonne ;
  - puces via le style de liste Word, jamais un caractere puce saisi a la
    main : un ATS reconnait la premiere, pas la seconde ;
  - styles Heading standards, pour que les sections soient detectees.

Prerequis : pip install python-docx lxml
"""

import pathlib
import sys

import lxml.html
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent))
from make_cv_pdf import INTERDITS, MOTS_CLES  # noqa: E402  source unique des regles

RACINE = pathlib.Path(__file__).resolve().parent.parent
SOURCE = RACINE / "cv" / "cv.html"
SORTIES = [RACINE / "cv" / "CV-Moussa-Zedira.docx", RACINE / "public" / "CV.MoussaZedira.docx"]

SERIF = "Georgia"      # tient le role de Newsreader
SANS = "Arial"         # tient le role d'IBM Plex Sans

# Memes marges que le PDF : la colonne de texte fait 182 mm dans les deux
# documents, donc les retours a la ligne tombent au meme endroit.
MARGE_V, MARGE_H = Mm(12), Mm(14)
LARGEUR_TEXTE = Mm(210 - 2 * 14)

# Le .docx n'a pas la contrainte de densite du PDF : a contenu egal il
# reste un tiers de page libre. On depense cette place en blancs plutot
# que de la laisser en bas de page.
INTERLIGNE = 1.12

ENCRE = RGBColor(0x16, 0x18, 0x1C)
TEXTE = RGBColor(0x33, 0x37, 0x3D)
GRIS = RGBColor(0x6B, 0x70, 0x79)
AMBRE = RGBColor(0xB4, 0x53, 0x09)


# ----------------------------------------------------------------------
# Lecture de la source
# ----------------------------------------------------------------------
def texte(element):
    return " ".join(element.text_content().split())


def lignes_contact(paragraphe):
    """Decoupe le bloc de contact sur les <br>, en gardant l'ordre."""
    resultat, courante = [], [paragraphe.text or ""]
    for enfant in paragraphe:
        if enfant.tag == "br":
            resultat.append(courante)
            courante = []
        else:
            courante.append(enfant.text_content())
        courante.append(enfant.tail or "")
    resultat.append(courante)
    return [" ".join("".join(bloc).split()) for bloc in resultat if "".join(bloc).strip()]


def lire_source():
    arbre = lxml.html.fromstring(SOURCE.read_text(encoding="utf-8"))
    page = arbre.find_class("page")[0]
    entete = arbre.find_class("header")[0]

    contenu = {
        "nom": texte(entete.find("h1")),
        "role": texte(entete.find_class("role")[0]),
        "objectif": texte(entete.find_class("objectif")[0]),
        "contact": lignes_contact(entete.find_class("contact")[0]),
        "sections": [],
    }

    for section in page.iter("section"):
        bloc = {"titre": texte(section.find("h2")), "competences": [], "entrees": []}

        for ligne in section.find_class("skills-row"):
            bloc["competences"].append(
                (texte(ligne.find_class("skills-label")[0]),
                 texte(ligne.find_class("skills-value")[0]))
            )

        for item in section.find_class("item"):
            dates = item.find_class("date")
            org = item.find_class("org")
            desc = item.find_class("desc")

            # Le lien vit dans le <h3> mais ne doit pas heriter du gras
            # serif de l'intitule : on le sort avant de lire le titre.
            titre_html = item.find_class("item-head")[0].find("h3")
            liens = titre_html.find_class("link")
            lien = texte(liens[0]) if liens else ""
            if liens:
                liens[0].drop_tree()

            bloc["entrees"].append({
                "titre": texte(titre_html),
                "lien": lien,
                "date": texte(dates[0]) if dates else "",
                "org": texte(org[0]) if org else "",
                "desc": texte(desc[0]) if desc else "",
                "puces": [texte(li) for li in item.iter("li")],
            })

        contenu["sections"].append(bloc)

    return contenu


# ----------------------------------------------------------------------
# Mise en forme
# ----------------------------------------------------------------------
def styler(run, police=SANS, taille=9, gras=False, couleur=TEXTE, interlettrage=None):
    run.font.name = police
    run.font.size = Pt(taille)
    run.bold = gras
    run.font.color.rgb = couleur
    if interlettrage:
        rPr = run._element.get_or_add_rPr()
        espacement = OxmlElement("w:spacing")
        espacement.set(qn("w:val"), str(interlettrage))  # vingtiemes de point
        rPr.append(espacement)
    return run


def para(doc, avant=0, apres=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(avant)
    p.paragraph_format.space_after = Pt(apres)
    p.paragraph_format.line_spacing = INTERLIGNE
    return p


def filet_bas(paragraphe, couleur="DCDAD6"):
    """Reproduit le filet sous les titres de section du PDF.

    Un filet de paragraphe, jamais un tableau d'une cellule : un tableau
    utilise comme separateur a une hauteur minimale et ressort comme une
    case vide, en plus d'etre mal parse par les ATS.
    """
    pPr = paragraphe._p.get_or_add_pPr()
    bordures = OxmlElement("w:pBdr")
    bas = OxmlElement("w:bottom")
    bas.set(qn("w:val"), "single")
    bas.set(qn("w:sz"), "4")
    bas.set(qn("w:space"), "2")
    bas.set(qn("w:color"), couleur)
    bordures.append(bas)
    pPr.insert_element_before(
        bordures, "w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr", "w:sectPr"
    )


def construire(contenu):
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = SANS
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = INTERLIGNE

    for section in doc.sections:
        section.page_width, section.page_height = Mm(210), Mm(297)
        section.top_margin = section.bottom_margin = MARGE_V
        section.left_margin = section.right_margin = MARGE_H

    # En-tete : dans le corps du document, jamais dans un en-tete Word.
    styler(para(doc, apres=2).add_run(contenu["nom"]),
           police=SERIF, taille=20, gras=True, couleur=ENCRE)
    styler(para(doc, apres=5).add_run(contenu["role"].upper()),
           taille=8, gras=True, couleur=AMBRE, interlettrage=36)
    styler(para(doc, apres=5).add_run(contenu["objectif"]), taille=9, couleur=ENCRE)
    for ligne in contenu["contact"]:
        styler(para(doc, apres=1.5).add_run(ligne), taille=8, couleur=GRIS)

    for bloc in contenu["sections"]:
        titre = doc.add_heading(level=1)
        titre.paragraph_format.space_before = Pt(17)
        titre.paragraph_format.space_after = Pt(7)
        styler(titre.add_run(bloc["titre"].upper()),
               taille=8, gras=True, couleur=GRIS, interlettrage=32)
        filet_bas(titre)

        for label, valeur in bloc["competences"]:
            p = para(doc, apres=4.5)
            styler(p.add_run(f"{label} : "), gras=True, couleur=ENCRE)
            styler(p.add_run(valeur))

        for entree in bloc["entrees"]:
            p = para(doc, avant=11, apres=1.5)
            styler(p.add_run(entree["titre"]), police=SERIF, taille=10.5,
                   gras=True, couleur=ENCRE)
            if entree["lien"]:
                styler(p.add_run(f"  {entree['lien']}"), taille=8, couleur=AMBRE)
            if entree["date"]:
                # Tabulation alignee a droite : reproduit la date en bout de
                # ligne du PDF sans recourir a un tableau.
                p.paragraph_format.tab_stops.add_tab_stop(
                    LARGEUR_TEXTE, WD_TAB_ALIGNMENT.RIGHT
                )
                styler(p.add_run(f"\t{entree['date']}"), taille=8, couleur=GRIS)
            if entree["org"]:
                styler(para(doc, apres=2).add_run(entree["org"]), couleur=GRIS)
            if entree["desc"]:
                styler(para(doc).add_run(entree["desc"]))
            for puce in entree["puces"]:
                # Style de liste Word : c'est lui que l'ATS reconnait comme
                # une puce. Un caractere puce tape a la main ne l'est pas.
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2.5)
                p.paragraph_format.line_spacing = INTERLIGNE
                p.paragraph_format.left_indent = Mm(4.5)
                p.paragraph_format.first_line_indent = Mm(-3.5)
                styler(p.add_run(puce))

    return doc


# ----------------------------------------------------------------------
# Controles
# ----------------------------------------------------------------------
def verifier(chemin):
    """Rejoue sur le .docx les controles ATS du PDF."""
    doc = Document(str(chemin))
    lu = "\n".join(p.text for p in doc.paragraphs)
    problemes = []

    manquants = [m for m in MOTS_CLES if m.lower() not in lu.lower()]
    if manquants:
        problemes.append(f"mots-cles absents : {manquants}")
    for caractere, explication in INTERDITS.items():
        if caractere in lu:
            problemes.append(f"{lu.count(caractere)}x {explication}")
    if doc.tables:
        problemes.append(f"{len(doc.tables)} tableau(x) : mal parses par les ATS")
    if doc.element.body.findall(".//{*}txbxContent"):
        problemes.append("zone(s) de texte : ignorees par les ATS")
    for zone, nom in ((doc.sections[0].header, "en-tete"), (doc.sections[0].footer, "pied de page")):
        if any(p.text.strip() for p in zone.paragraphs):
            problemes.append(f"{nom} Word non vide : son contenu sera ignore par les ATS")

    return problemes, lu


def compter_pages(chemin):
    """Pagination reelle via Word. Retourne None si Word est indisponible.

    python-docx ne sait pas paginer : seul un moteur de rendu le peut. On
    ne fait donc pas echouer la generation faute de Word, on le signale.
    """
    try:
        import win32com.client
    except ImportError:
        return None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        document = word.Documents.Open(str(chemin), False, True)
        pages = document.ComputeStatistics(2)  # wdStatisticPages
        document.Close(False)
        word.Quit()
        return pages
    except Exception:
        return None


def main():
    construire(lire_source()).save(SORTIES[0])

    problemes, lu = verifier(SORTIES[0])
    pages = compter_pages(SORTIES[0])

    taille = SORTIES[0].stat().st_size / 1024
    etat_pages = f"{pages} page(s)" if pages else "pagination non verifiee (Word absent)"
    print(f"{SORTIES[0].relative_to(RACINE)} : {taille:.0f} Ko | {len(lu)} caracteres | {etat_pages}")

    if pages and pages > 1:
        problemes.append(f"{pages} pages : le CV doit tenir sur une seule")

    if problemes:
        print("\nREGRESSION ATS :", file=sys.stderr)
        for probleme in problemes:
            print(f"  - {probleme}", file=sys.stderr)
        return 1

    for copie in SORTIES[1:]:
        copie.write_bytes(SORTIES[0].read_bytes())
        print(f"{copie.relative_to(RACINE)} : copie")

    print(f"\nOK - {len(MOTS_CLES)} mots-cles presents, aucun caractere hostile.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
